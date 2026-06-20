"""
JOBSTER — Agent IA
===================
Cet agent utilise Ollama + Qwen pour :
  1. Chercher des offres en temps réel (France Travail + Adzuna)
  2. Analyser une offre depuis une URL
  3. Calculer un score de matching CV / offre
  4. Générer une lettre de motivation personnalisée en .docx

Comment utiliser :
  1. Lance Ollama en arrière-plan (il tourne déjà si installé)
  2. Ouvre un terminal dans le dossier scraping
  3. Tape : python jobster_agent.py
"""

import json
import subprocess
import os
import sys
import re
import requests
from bs4 import BeautifulSoup

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import llm_backend as ollama  # bascule Ollama (local) / Groq (cloud) — voir scraping/llm_backend.py

# ============================================================
# CONFIGURATION DE L'AGENT
# On utilise le modèle Qwen3 léger installé sur ton PC.
# Il tourne entièrement en local — pas besoin d'internet.
# ============================================================

MODELE = "qwen3:1.7b"


# ============================================================
# OUTIL 2 — ANALYSEUR D'OFFRE DEPUIS UNE URL
# L'utilisateur colle un lien d'offre d'emploi.
# On va sur la page, on lit le contenu, et Qwen extrait
# toutes les informations utiles de façon structurée.
# ============================================================

def analyser_offre_url(url):
    """
    Visite une page d'offre d'emploi et extrait les informations
    clés : titre, entreprise, missions, compétences, salaire.
    """
    print(f"\n  [Outil 2] Analyse de l'offre : {url}")

    # Etape 1 : on récupère le contenu de la page
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/120.0.0.0 Safari/537.36",
            "Accept-Language": "fr-FR,fr;q=0.9",
        }
        reponse = requests.get(url, headers=headers, timeout=15)
        soup = BeautifulSoup(reponse.text, "html.parser")

        # On supprime les balises script et style qui pollueraient le texte
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        # On extrait le texte brut de la page
        texte_brut = soup.get_text(separator="\n", strip=True)

        # On limite à 3000 caractères pour ne pas dépasser la mémoire de Qwen
        texte_limite = texte_brut[:3000]
        print(f"  [Outil 2] Page récupérée — {len(texte_brut)} caractères")

    except Exception as e:
        return f"Je n'ai pas pu accéder à cette page : {e}"

    # Etape 2 : Qwen analyse le texte et extrait les infos
    prompt = f"""Tu es un expert RH. Analyse ce texte d'une offre d'emploi et extrais les informations clés.

Texte de l'offre :
{texte_limite}

Réponds en français avec exactement cette structure :
**Titre du poste :** ...
**Entreprise :** ...
**Lieu :** ...
**Type de contrat :** ...
**Salaire :** ... (si mentionné, sinon "Non précisé")
**Missions principales :** liste des 3-5 missions clés
**Compétences requises :** liste des compétences techniques et soft skills
**Niveau d'expérience :** ...
**Points forts de l'offre :** 2-3 éléments attractifs
"""

    try:
        reponse_ia = ollama.chat(
            model=MODELE,
            messages=[{"role": "user", "content": prompt}],
            options={"temperature": 0.3}
        )
        return reponse_ia["message"]["content"].strip()
    except Exception as e:
        return f"Erreur lors de l'analyse : {e}"


# ============================================================
# OUTIL 3 — ANALYSEUR DE MATCHING CV / OFFRE
# L'utilisateur peut :
#   - Coller son CV en texte directement dans le chat
#   - Donner le chemin d'un fichier PDF sur son PC
#   - Utiliser profil.json s'il existe déjà
# Jobster compare le CV avec l'offre et donne un score /100.
# Le profil évolue à chaque fois — pas de données figées.
# ============================================================

FICHIER_PROFIL = os.path.join(os.path.dirname(os.path.abspath(__file__)), "profil.json")


def lire_pdf(chemin_fichier):
    """
    Lit le contenu texte d'un fichier PDF.
    Nécessite : pip install pdfplumber
    """
    try:
        import pdfplumber
        texte = ""
        with pdfplumber.open(chemin_fichier) as pdf:
            for page in pdf.pages:
                texte += page.extract_text() or ""
        return texte.strip()
    except ImportError:
        return None
    except Exception as e:
        return None


def extraire_cv_depuis_message(message):
    """
    Essaie de trouver le CV dans le message de l'utilisateur.
    Retourne (texte_cv, url_offre, mode_cv).
    mode_cv : 'texte', 'pdf', 'profil', ou None
    """
    # Cherche une URL d'offre dans le message
    urls = re.findall(r'https?://[^\s]+', message)
    url_offre = urls[0] if urls else None

    # Cherche un chemin de fichier PDF dans le message
    # Exemple : C:\Users\gilda\Documents\mon_cv.pdf
    chemins_pdf = re.findall(r'[A-Za-z]:\\[^\s]+\.pdf|/[^\s]+\.pdf', message, re.IGNORECASE)
    if chemins_pdf:
        chemin = chemins_pdf[0]
        texte_pdf = lire_pdf(chemin)
        if texte_pdf:
            return texte_pdf, url_offre, "pdf"

    # Cherche si l'utilisateur a collé du texte de CV
    # On détecte des mots clés typiques d'un CV
    mots_cv = ["expérience", "formation", "compétences", "diplôme",
                "master", "licence", "bac", "stage", "cdi", "cdd",
                "skills", "education", "experience", "curriculum"]
    message_lower = message.lower()
    nb_mots_cv = sum(1 for mot in mots_cv if mot in message_lower)

    if nb_mots_cv >= 3:
        # Le message contient probablement un CV collé
        # On enlève l'URL du texte pour garder seulement le CV
        texte_cv = message
        for url in urls:
            texte_cv = texte_cv.replace(url, "").strip()
        return texte_cv, url_offre, "texte"

    # Fallback : utilise profil.json si il existe
    if os.path.exists(FICHIER_PROFIL):
        with open(FICHIER_PROFIL, "r", encoding="utf-8") as f:
            profil = json.load(f)
        return json.dumps(profil, ensure_ascii=False), url_offre, "profil"

    return None, url_offre, None


def calculer_matching(message_complet):
    """
    Compare le CV de l'utilisateur avec une offre d'emploi
    et calcule un score de compatibilité sur 100 avec explications.
    Supporte : texte collé, PDF, ou profil.json
    """
    print(f"\n  [Outil 3] Calcul du score de matching...")

    # Etape 1 : extraire le CV et l'URL depuis le message
    texte_cv, url_offre, mode = extraire_cv_depuis_message(message_complet)

    if not url_offre:
        return "Pour calculer le matching, colle aussi le lien de l'offre dans ton message.\nExemple : 'score mon CV : [texte CV] avec https://...'"

    if not texte_cv:
        return ("Je n'ai pas trouvé ton CV. Tu peux :\n"
                "1. Coller le texte de ton CV directement\n"
                "2. Donner le chemin de ton PDF : C:\\Users\\...\\cv.pdf\n"
                "3. Créer un fichier profil.json dans le dossier scraping")

    if mode == "pdf":
        print(f"  [Outil 3] CV lu depuis fichier PDF")
    elif mode == "texte":
        print(f"  [Outil 3] CV détecté dans le message ({len(texte_cv)} caractères)")
    elif mode == "profil":
        print(f"  [Outil 3] Profil chargé depuis profil.json")

    # Etape 2 : récupérer le contenu de l'offre
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/120.0.0.0 Safari/537.36",
        }
        reponse = requests.get(url_offre, headers=headers, timeout=15)
        soup = BeautifulSoup(reponse.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer"]):
            tag.decompose()
        texte_offre = soup.get_text(separator="\n", strip=True)[:2500]
        print(f"  [Outil 3] Offre récupérée — {len(texte_offre)} caractères")
    except Exception as e:
        return f"Je n'ai pas pu accéder à cette offre : {e}"

    # Etape 3 : Qwen calcule le score
    # On limite le CV à 2000 caractères pour ne pas surcharger Qwen
    texte_cv_limite = texte_cv[:2000]

    prompt = f"""Tu es un expert RH et recruteur senior.
Compare ce CV avec cette offre d'emploi et donne un score de matching précis.

CV DU CANDIDAT :
{texte_cv_limite}

OFFRE D'EMPLOI :
{texte_offre}

Réponds en français avec exactement cette structure :

🎯 SCORE DE MATCHING : XX/100

✅ POINTS FORTS (ce qui correspond bien) :
- point 1
- point 2
- point 3

⚠️ POINTS DE VIGILANCE (ce qui manque ou diverge) :
- point 1
- point 2

💡 CONSEIL PERSONNALISÉ :
Un conseil concret et actionnable pour maximiser les chances.

📌 VERDICT :
Une phrase de conclusion franche : faut-il postuler ou non ?
"""

    try:
        reponse_ia = ollama.chat(
            model=MODELE,
            messages=[{"role": "user", "content": prompt}],
            options={"temperature": 0.4}
        )
        return reponse_ia["message"]["content"].strip()
    except Exception as e:
        return f"Erreur lors du calcul de matching : {e}"


# ============================================================
# OUTIL 4 — GENERATEUR DE LETTRE DE MOTIVATION (.docx)
# L'utilisateur donne son CV + le lien de l'offre.
# Qwen rédige une lettre personnalisée, et on la sauvegarde
# en fichier Word (.docx) prêt à télécharger et envoyer.
# La lettre est adaptée à l'offre réelle — pas un modèle générique.
# ============================================================

def generer_lettre_motivation(message_complet):
    """
    Génère une lettre de motivation personnalisée en .docx
    à partir du CV de l'utilisateur et d'une offre d'emploi.
    """
    print(f"\n  [Outil 4] Génération de la lettre de motivation...")

    # Etape 1 : extraire le CV et l'URL
    texte_cv, url_offre, mode = extraire_cv_depuis_message(message_complet)

    if not url_offre:
        return "Pour générer une lettre, colle aussi le lien de l'offre.\nEx: 'lettre [ton CV] https://...'"

    if not texte_cv:
        return "Je n'ai pas trouvé ton CV. Colle ton CV avec le lien de l'offre."

    # Etape 2 : récupérer le contenu de l'offre
    try:
        headers = {"User-Agent": "Mozilla/5.0 Chrome/120.0.0.0 Safari/537.36"}
        reponse = requests.get(url_offre, headers=headers, timeout=15)
        soup = BeautifulSoup(reponse.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer"]):
            tag.decompose()
        texte_offre = soup.get_text(separator="\n", strip=True)[:2000]
        print(f"  [Outil 4] Offre récupérée — {len(texte_offre)} caractères")
    except Exception as e:
        return f"Je n'ai pas pu accéder à l'offre : {e}"

    # Etape 3 : Qwen rédige la lettre
    prompt = f"""Tu es un expert en recrutement et rédaction professionnelle.
Rédige une lettre de motivation professionnelle en français, personnalisée et convaincante.

CV DU CANDIDAT :
{texte_cv[:1500]}

OFFRE D'EMPLOI :
{texte_offre}

Rédige une lettre de motivation avec exactement cette structure :
- En-tête : Nom, coordonnées du candidat (déduits du CV)
- Date d'aujourd'hui
- Objet : Candidature pour [poste] chez [entreprise]
- Corps en 3 paragraphes :
  1. Accroche percutante qui montre la connaissance de l'entreprise
  2. Pourquoi le candidat est fait pour ce poste (expériences concrètes)
  3. Projection et motivation
- Formule de politesse professionnelle
- Signature

Sois précis, professionnel et personnalisé. Pas de formules génériques.
"""

    try:
        reponse_ia = ollama.chat(
            model=MODELE,
            messages=[{"role": "user", "content": prompt}],
            options={"temperature": 0.7}
        )
        contenu_lettre = reponse_ia["message"]["content"].strip()
    except Exception as e:
        return f"Erreur lors de la rédaction : {e}"

    # Etape 4 : sauvegarder en .docx
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import datetime

        doc = DocxDocument()

        # Marges
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Contenu de la lettre paragraphe par paragraphe
        for ligne in contenu_lettre.split("\n"):
            p = doc.add_paragraph()
            run = p.add_run(ligne)
            run.font.name = "Arial"
            run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(6)

        # Nom du fichier avec la date
        date_str = datetime.datetime.now().strftime("%Y%m%d_%H%M")
        nom_fichier = f"lettre_motivation_jobster_{date_str}.docx"
        chemin_fichier = os.path.join(os.path.dirname(os.path.abspath(__file__)), nom_fichier)
        doc.save(chemin_fichier)

        print(f"  [Outil 4] Lettre sauvegardée : {nom_fichier}")
        return f"✅ Lettre de motivation générée !\n\nFichier : {chemin_fichier}\n\n--- APERÇU ---\n\n{contenu_lettre[:500]}..."

    except ImportError:
        # Si python-docx n'est pas installé, on retourne juste le texte
        print("  [Outil 4] python-docx non installé — affichage texte uniquement")
        return f"📄 LETTRE DE MOTIVATION :\n\n{contenu_lettre}\n\n(Pour sauvegarder en .docx : pip install python-docx)"
    except Exception as e:
        return f"Lettre rédigée mais erreur de sauvegarde : {e}\n\n{contenu_lettre}"


# ============================================================
# OUTIL 5 — GENERATEUR DE CV ADAPTE EN PDF
# L'utilisateur donne son CV + le lien de l'offre.
# Qwen réécrit le CV en mettant en avant les expériences
# les plus pertinentes et en intégrant les mots-clés de
# l'offre. Le résultat est sauvegardé en PDF professionnel.
# ============================================================

def generer_cv_adapte(message_complet):
    """
    Génère un CV adapté à une offre spécifique en PDF.
    Le CV est réécrit pour coller exactement au poste visé.
    """
    print(f"\n  [Outil 5] Génération du CV adapté en PDF...")

    # Etape 1 : extraire le CV et l'URL
    texte_cv, url_offre, mode = extraire_cv_depuis_message(message_complet)

    if not url_offre:
        return "Pour adapter ton CV, colle aussi le lien de l'offre.\nEx: 'cv [ton CV] https://...'"

    if not texte_cv:
        return "Je n'ai pas trouvé ton CV. Colle ton CV avec le lien de l'offre."

    # Etape 2 : récupérer le contenu de l'offre
    try:
        headers = {"User-Agent": "Mozilla/5.0 Chrome/120.0.0.0 Safari/537.36"}
        reponse = requests.get(url_offre, headers=headers, timeout=15)
        soup = BeautifulSoup(reponse.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer"]):
            tag.decompose()
        texte_offre = soup.get_text(separator="\n", strip=True)[:2000]
        print(f"  [Outil 5] Offre récupérée — {len(texte_offre)} caractères")
    except Exception as e:
        return f"Je n'ai pas pu accéder à l'offre : {e}"

    # Etape 3 : Qwen adapte le CV
    prompt = f"""Tu es un expert en recrutement et optimisation de CV.
Adapte ce CV pour qu'il corresponde parfaitement à cette offre d'emploi.

CV ORIGINAL :
{texte_cv[:1500]}

OFFRE D'EMPLOI :
{texte_offre}

Réécris le CV adapté avec exactement ces sections :
**NOM ET COORDONNEES** : (reprends du CV original)

**PROFIL** : 3 phrases percutantes qui correspondent exactement au poste

**EXPERIENCES PROFESSIONNELLES** : (réordonne et reformule pour mettre en avant ce qui est pertinent pour ce poste, avec les mots-clés de l'offre)

**FORMATIONS** : (reprends du CV original)

**COMPETENCES CLES** : (sélectionne et met en avant les compétences qui matchent l'offre)

**LANGUES** : (reprends du CV original)

Intègre naturellement les mots-clés importants de l'offre dans le CV.
Sois concis et percutant. Format professionnel.
"""

    try:
        reponse_ia = ollama.chat(
            model=MODELE,
            messages=[{"role": "user", "content": prompt}],
            options={"temperature": 0.5}
        )
        contenu_cv = reponse_ia["message"]["content"].strip()
    except Exception as e:
        return f"Erreur lors de l'adaptation du CV : {e}"

    # Etape 4 : générer le PDF avec reportlab
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        import datetime

        date_str = datetime.datetime.now().strftime("%Y%m%d_%H%M")
        nom_fichier = f"cv_adapte_jobster_{date_str}.pdf"
        chemin_fichier = os.path.join(os.path.dirname(os.path.abspath(__file__)), nom_fichier)

        doc = SimpleDocTemplate(
            chemin_fichier,
            pagesize=A4,
            topMargin=1.5*cm,
            bottomMargin=1.5*cm,
            leftMargin=2*cm,
            rightMargin=2*cm
        )

        styles = getSampleStyleSheet()

        # Style titre principal (nom)
        style_nom = ParagraphStyle(
            "Nom",
            parent=styles["Title"],
            fontSize=18,
            textColor=colors.HexColor("#1a1a2e"),
            spaceAfter=4,
            fontName="Helvetica-Bold"
        )

        # Style section
        style_section = ParagraphStyle(
            "Section",
            parent=styles["Heading2"],
            fontSize=11,
            textColor=colors.HexColor("#16213e"),
            spaceBefore=10,
            spaceAfter=4,
            fontName="Helvetica-Bold",
            borderPad=2,
        )

        # Style corps
        style_corps = ParagraphStyle(
            "Corps",
            parent=styles["Normal"],
            fontSize=9.5,
            leading=14,
            spaceAfter=3,
            fontName="Helvetica"
        )

        story = []

        # On parse le contenu généré par Qwen
        lignes = contenu_cv.split("\n")
        premiere_ligne = True

        for ligne in lignes:
            ligne = ligne.strip()
            if not ligne:
                story.append(Spacer(1, 4))
                continue

            # Titres de section (commencent par **)
            if ligne.startswith("**") and ligne.endswith("**"):
                titre = ligne.replace("**", "").strip()
                if premiere_ligne:
                    story.append(Paragraph(titre, style_nom))
                    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#16213e")))
                    premiere_ligne = False
                else:
                    story.append(Paragraph(titre, style_section))
                    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc")))
            else:
                # Nettoyer les marqueurs markdown restants
                ligne = ligne.replace("**", "").replace("*", "").replace("##", "").replace("#", "")
                if ligne:
                    story.append(Paragraph(ligne, style_corps))

        doc.build(story)
        print(f"  [Outil 5] CV PDF sauvegardé : {nom_fichier}")
        return (f"✅ CV adapté généré en PDF !\n\n"
                f"Fichier : {chemin_fichier}\n\n"
                f"--- APERÇU DU CV ADAPTÉ ---\n\n"
                f"{contenu_cv[:600]}...")

    except ImportError:
        print("  [Outil 5] reportlab non installé — affichage texte uniquement")
        return (f"📄 CV ADAPTÉ :\n\n{contenu_cv}\n\n"
                f"(Pour sauvegarder en PDF : pip install reportlab)")
    except Exception as e:
        return f"CV rédigé mais erreur PDF : {e}\n\n{contenu_cv}"


# ============================================================
# OUTIL 6 — GENERATEUR DE COPIER-COLLER ET MAILS
# L'utilisateur donne son CV + le lien de l'offre.
# Jobster génère :
#   1. Un texte de présentation pour les formulaires en ligne
#   2. Un mail de candidature prêt à envoyer
#   3. Un message LinkedIn pour contacter le recruteur
# Ces textes sont sauvegardés dans un fichier .txt
# ============================================================

def generer_copier_coller(message_complet):
    """
    Génère des textes prêts à l'emploi pour les candidatures :
    - Texte formulaire (copier-coller)
    - Mail de candidature complet
    - Message LinkedIn pour le recruteur
    """
    print(f"\n  [Outil 6] Génération des textes de candidature...")

    # Etape 1 : extraire le CV et l'URL
    texte_cv, url_offre, mode = extraire_cv_depuis_message(message_complet)

    if not url_offre:
        return "Pour générer les textes, colle aussi le lien de l'offre.\nEx: 'mail [CV] https://...'"

    if not texte_cv:
        return "Je n'ai pas trouvé ton CV. Colle ton CV avec le lien de l'offre."

    # Etape 2 : récupérer le contenu de l'offre
    try:
        headers = {"User-Agent": "Mozilla/5.0 Chrome/120.0.0.0 Safari/537.36"}
        reponse = requests.get(url_offre, headers=headers, timeout=15)
        soup = BeautifulSoup(reponse.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer"]):
            tag.decompose()
        texte_offre = soup.get_text(separator="\n", strip=True)[:2000]
        print(f"  [Outil 6] Offre récupérée — {len(texte_offre)} caractères")
    except Exception as e:
        return f"Je n'ai pas pu accéder à l'offre : {e}"

    # Etape 3 : Qwen génère les 3 textes
    prompt = f"""Tu es un expert en recrutement et communication professionnelle.
Génère 3 textes de candidature personnalisés en français pour ce profil et cette offre.

CV DU CANDIDAT :
{texte_cv[:1200]}

OFFRE D'EMPLOI :
{texte_offre}

Génère exactement ces 3 textes séparés par ---SEPARATEUR--- :

TEXTE 1 — FORMULAIRE EN LIGNE (300 mots max)
Texte de présentation à coller dans les champs "Parlez-nous de vous" ou "Lettre de motivation courte" des formulaires en ligne. Percutant, professionnel, adapté à l'offre.

---SEPARATEUR---

TEXTE 2 — MAIL DE CANDIDATURE COMPLET
Objet : [objet accrocheur]
Corps du mail complet avec bonjour, présentation, motivation, appel à action et signature professionnelle.

---SEPARATEUR---

TEXTE 3 — MESSAGE LINKEDIN (150 mots max)
Message court et direct pour contacter le recruteur sur LinkedIn. Professionnel mais humain.
"""

    try:
        reponse_ia = ollama.chat(
            model=MODELE,
            messages=[{"role": "user", "content": prompt}],
            options={"temperature": 0.7}
        )
        contenu = reponse_ia["message"]["content"].strip()
    except Exception as e:
        return f"Erreur lors de la génération : {e}"

    # Etape 4 : sauvegarder dans un fichier .txt
    try:
        import datetime
        date_str = datetime.datetime.now().strftime("%Y%m%d_%H%M")
        nom_fichier = f"candidature_jobster_{date_str}.txt"
        chemin_fichier = os.path.join(os.path.dirname(os.path.abspath(__file__)), nom_fichier)

        # On structure le fichier proprement
        parties = contenu.split("---SEPARATEUR---")
        contenu_fichier = "=" * 60 + "\n"
        contenu_fichier += "JOBSTER — TEXTES DE CANDIDATURE\n"
        contenu_fichier += "=" * 60 + "\n\n"

        titres = ["📋 TEXTE FORMULAIRE EN LIGNE", "📧 MAIL DE CANDIDATURE", "💼 MESSAGE LINKEDIN"]
        for i, (titre, partie) in enumerate(zip(titres, parties)):
            contenu_fichier += titre + "\n"
            contenu_fichier += "-" * 40 + "\n"
            contenu_fichier += partie.strip() + "\n\n"
            contenu_fichier += "=" * 60 + "\n\n"

        with open(chemin_fichier, "w", encoding="utf-8") as f:
            f.write(contenu_fichier)

        print(f"  [Outil 6] Textes sauvegardés : {nom_fichier}")

        # Aperçu dans le terminal
        apercu = parties[0].strip()[:400] if parties else contenu[:400]
        return (f"✅ Textes de candidature générés !\n\n"
                f"Fichier : {chemin_fichier}\n\n"
                f"--- APERÇU (Texte formulaire) ---\n\n"
                f"{apercu}...")

    except Exception as e:
        return f"Textes générés mais erreur de sauvegarde : {e}\n\n{contenu[:600]}"


# ============================================================
# OUTIL 7 — GENERATEUR D'URLS INTELLIGENTES
# À partir d'une offre d'emploi, Jobster génère des liens
# directs utiles pour le candidat :
#   - Lien vers le formulaire de candidature
#   - Recherche LinkedIn du recruteur/RH de l'entreprise
#   - Page Glassdoor de l'entreprise (avis, salaires)
#   - Google News sur l'entreprise (actualités récentes)
#   - Page Société.com (données légales)
# ============================================================

def generer_urls_intelligentes(url_offre):
    """
    Analyse une offre d'emploi et génère des URLs utiles
    pour le candidat : recruteur LinkedIn, Glassdoor, etc.
    """
    print(f"\n  [Outil 7] Génération des URLs intelligentes...")

    # Etape 1 : récupérer le contenu de l'offre
    try:
        headers = {"User-Agent": "Mozilla/5.0 Chrome/120.0.0.0 Safari/537.36"}
        reponse = requests.get(url_offre, headers=headers, timeout=15)
        soup = BeautifulSoup(reponse.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer"]):
            tag.decompose()
        texte_offre = soup.get_text(separator="\n", strip=True)[:2000]
        print(f"  [Outil 7] Offre récupérée — {len(texte_offre)} caractères")
    except Exception as e:
        return f"Je n'ai pas pu accéder à l'offre : {e}"

    # Etape 2 : Qwen extrait le nom de l'entreprise et le poste
    prompt_extraction = f"""Analyse cette offre d'emploi.
Réponds UNIQUEMENT avec un JSON valide sans texte autour :

{texte_offre[:1000]}

{{"entreprise": "nom exact de l'entreprise", "poste": "intitulé exact du poste", "ville": "ville du poste"}}
"""
    try:
        reponse_ia = ollama.chat(
            model=MODELE,
            messages=[{"role": "user", "content": prompt_extraction}],
            options={"temperature": 0}
        )
        texte = reponse_ia["message"]["content"].strip()
        debut = texte.find("{")
        fin = texte.rfind("}") + 1
        infos = json.loads(texte[debut:fin]) if debut != -1 else {}
        entreprise = infos.get("entreprise", "").strip()
        poste = infos.get("poste", "").strip()
    except Exception:
        entreprise = ""
        poste = ""

    print(f"  [Outil 7] Entreprise : {entreprise} | Poste : {poste}")

    # Etape 3 : construire les URLs intelligentes
    entreprise_encode = entreprise.replace(" ", "%20")
    entreprise_plus = entreprise.replace(" ", "+")

    url_linkedin_recruteur = (
        f"https://www.linkedin.com/search/results/people/?"
        f"keywords={entreprise_encode}%20recruteur%20RH&origin=GLOBAL_SEARCH_HEADER"
    )
    url_linkedin_entreprise = (
        f"https://www.linkedin.com/company/{entreprise.lower().replace(' ', '-')}"
    )
    url_glassdoor = f"https://www.glassdoor.fr/Avis/{entreprise_plus}-Avis-E0.htm"
    url_news = f"https://news.google.com/search?q={entreprise_encode}&hl=fr&gl=FR&ceid=FR:fr"
    url_societe = f"https://www.societe.com/cgi-bin/search?champs={entreprise_encode}"
    url_indeed = f"https://fr.indeed.com/cmp/{entreprise_plus}/jobs"

    resultat = f"🔗 LIENS UTILES POUR CANDIDATER CHEZ {entreprise.upper()} :\n\n"
    resultat += f"📋 POSTULER\n   {url_offre}\n\n"
    resultat += f"👤 RECRUTEUR LINKEDIN\n   {url_linkedin_recruteur}\n"
    resultat += f"   Page entreprise : {url_linkedin_entreprise}\n\n"
    resultat += f"⭐ GLASSDOOR (avis & salaires)\n   {url_glassdoor}\n\n"
    resultat += f"📰 ACTUALITÉS GOOGLE NEWS\n   {url_news}\n\n"
    resultat += f"🏛️  DONNÉES LÉGALES (Société.com)\n   {url_societe}\n\n"
    resultat += f"💼 AUTRES OFFRES DE {entreprise.upper()}\n   {url_indeed}\n\n"
    resultat += "💡 Commence par Glassdoor pour la culture d'entreprise,\n"
    resultat += "   puis LinkedIn pour contacter directement le recruteur !"

    # Sauvegarder dans un fichier .txt
    try:
        import datetime
        date_str = datetime.datetime.now().strftime("%Y%m%d_%H%M")
        nom_fichier = f"liens_{entreprise.replace(' ', '_')}_{date_str}.txt"
        chemin_fichier = os.path.join(os.path.dirname(os.path.abspath(__file__)), nom_fichier)
        with open(chemin_fichier, "w", encoding="utf-8") as f:
            f.write(resultat)
        print(f"  [Outil 7] Liens sauvegardés : {nom_fichier}")
        return resultat + f"\n\n📁 Fichier : {chemin_fichier}"
    except Exception:
        return resultat


# ============================================================
# OUTIL 8 — SCRAPER D'INFOS ENTREPRISES
# L'utilisateur donne le nom d'une entreprise ou une URL d'offre.
# Jobster scrape automatiquement :
#   - Les avis employés sur Glassdoor
#   - Les actualités récentes via Google News
#   - Les données légales sur Société.com
# Tout est synthétisé par Qwen en un rapport clair.
# ============================================================

def scraper_infos_entreprise(message_complet):
    """Récupère et synthétise les infos sur une entreprise."""
    print(f"\n  [Outil 8] Scraping des infos entreprise...")

    urls_trouvees = re.findall(r'https?://[^\s]+', message_complet)
    nom_entreprise = ""

    if urls_trouvees:
        try:
            headers = {"User-Agent": "Mozilla/5.0 Chrome/120.0.0.0 Safari/537.36"}
            r = requests.get(urls_trouvees[0], headers=headers, timeout=15)
            soup = BeautifulSoup(r.text, "html.parser")
            for tag in soup(["script", "style"]): tag.decompose()
            texte = soup.get_text(separator="\n", strip=True)[:800]
            prompt = f"Extrait uniquement le nom de l'entreprise. Reponds UNIQUEMENT avec le nom.\n{texte}"
            rep = ollama.chat(model=MODELE, messages=[{"role": "user", "content": prompt}], options={"temperature": 0})
            nom_entreprise = rep["message"]["content"].strip().split("\n")[0]
        except:
            pass
    else:
        mots_a_enlever = ["entreprise", "infos sur", "informations sur", "que sais-tu de", "parle moi de", "recherche infos"]
        nom_entreprise = message_complet
        for mot in mots_a_enlever:
            nom_entreprise = nom_entreprise.lower().replace(mot, "").strip()

    if not nom_entreprise:
        return "Donne-moi le nom d'une entreprise.\nEx: 'entreprise Capgemini'"

    print(f"  [Outil 8] Entreprise cible : {nom_entreprise}")

    entreprise_plus = nom_entreprise.replace(" ", "+")
    entreprise_encode = nom_entreprise.replace(" ", "%20")

    infos_glassdoor = "Non accessible."
    infos_news = "Non accessible."
    infos_societe = "Non accessible."

    try:
        url = f"https://www.glassdoor.fr/Recherche/results.htm?keyword={entreprise_plus}"
        r = requests.get(url, headers={"User-Agent": "Mozilla/5.0 Chrome/120.0.0.0"}, timeout=10)
        soup = BeautifulSoup(r.text, "html.parser")
        for tag in soup(["script", "style"]): tag.decompose()
        infos_glassdoor = soup.get_text(separator=" ", strip=True)[:1200]
        print(f"  [Outil 8] Glassdoor OK")
    except: pass

    try:
        url = f"https://news.google.com/search?q={entreprise_encode}&hl=fr&gl=FR&ceid=FR:fr"
        r = requests.get(url, headers={"User-Agent": "Mozilla/5.0 Chrome/120.0.0.0"}, timeout=10)
        soup = BeautifulSoup(r.text, "html.parser")
        for tag in soup(["script", "style"]): tag.decompose()
        infos_news = soup.get_text(separator=" ", strip=True)[:1200]
        print(f"  [Outil 8] Google News OK")
    except: pass

    try:
        url = f"https://www.societe.com/cgi-bin/search?champs={entreprise_encode}"
        r = requests.get(url, headers={"User-Agent": "Mozilla/5.0 Chrome/120.0.0.0"}, timeout=10)
        soup = BeautifulSoup(r.text, "html.parser")
        for tag in soup(["script", "style"]): tag.decompose()
        infos_societe = soup.get_text(separator=" ", strip=True)[:1200]
        print(f"  [Outil 8] Societe.com OK")
    except: pass

    prompt = f"""Synthétise ces infos sur "{nom_entreprise}" pour un candidat.

GLASSDOOR : {infos_glassdoor[:500]}
ACTUALITES : {infos_news[:500]}
DONNEES LEGALES : {infos_societe[:500]}

Redige un rapport en français avec :
🏢 PRESENTATION DE L'ENTREPRISE
⭐ CULTURE ET AMBIANCE
📰 ACTUALITES CLES
🏛 DONNEES LEGALES
💡 3 CONSEILS POUR L'ENTRETIEN
"""

    try:
        rep = ollama.chat(model=MODELE, messages=[{"role": "user", "content": prompt}], options={"temperature": 0.4})
        rapport = rep["message"]["content"].strip()
    except Exception as e:
        return f"Erreur synthese : {e}"

    try:
        import datetime
        date_str = datetime.datetime.now().strftime("%Y%m%d_%H%M")
        nom_fichier = f"rapport_{nom_entreprise.replace(' ', '_')}_{date_str}.txt"
        chemin_fichier = os.path.join(os.path.dirname(os.path.abspath(__file__)), nom_fichier)
        with open(chemin_fichier, "w", encoding="utf-8") as f:
            f.write(f"RAPPORT JOBSTER\n{'='*60}\n\n{rapport}")
        print(f"  [Outil 8] Rapport sauvegarde : {nom_fichier}")
        return f"📊 RAPPORT — {nom_entreprise.upper()} :\n\n{rapport}\n\n📁 Fichier : {chemin_fichier}"
    except:
        return f"📊 RAPPORT — {nom_entreprise.upper()} :\n\n{rapport}"



# ============================================================
# OUTIL 9 — TRACKER DE CANDIDATURES (SQLite)
# Sauvegarde chaque candidature dans une base de données locale.
# L utilisateur peut ajouter, voir et mettre à jour ses candidatures.
# Statuts : envoyé, en attente, entretien, refus, accepté
# ============================================================

import sqlite3

DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "candidatures_jobster.db")

def init_db():
    """No-op — la DB est créée et migrée par server.py get_db()."""
    pass

def tracker_candidature(message_complet):
    """
    Gère le tracker de candidatures :
    - "tracker ajouter [entreprise] [poste]" → ajoute une candidature
    - "tracker voir" → affiche toutes les candidatures
    - "tracker statut [id] [nouveau statut]" → met à jour le statut
    """
    print(f"\n  [Outil 9] Tracker de candidatures...")
    init_db()
    d = message_complet.lower()
    import datetime

    # ── Voir toutes les candidatures
    if "voir" in d or "liste" in d or "mes candidatures" in d:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("SELECT id, date_ajout, entreprise, poste, statut, date_next_action FROM candidatures ORDER BY id DESC")
        rows = c.fetchall()
        conn.close()
        if not rows:
            return "Aucune candidature enregistrée.\nAjoute-en une avec : tracker ajouter [entreprise] [poste] [lien]"
        resultat = f"📋 TES CANDIDATURES ({len(rows)}) :\n\n"
        statuts_emoji = {"envoyé": "📤", "en attente": "⏳", "entretien": "🎯", "refus": "❌", "accepté": "🎉"}
        for row in rows:
            emoji = statuts_emoji.get(row[4], "📌")
            resultat += f"{emoji} [{row[0]}] {row[2]} — {row[3]}\n"
            resultat += f"   📅 Envoyé le {row[1]} | Statut : {row[4]}\n"
            if row[5]:
                resultat += f"   🔔 Relance prévue : {row[5]}\n"
            resultat += "\n"
        return resultat

    # ── Mettre à jour le statut
    elif "statut" in d or "mettre à jour" in d or "update" in d:
        # Format attendu : "tracker statut 3 entretien"
        mots = message_complet.split()
        statuts_valides = ["envoyé", "en attente", "entretien", "refus", "accepté"]
        id_trouve = None
        nouveau_statut = None
        for mot in mots:
            if mot.isdigit():
                id_trouve = int(mot)
            if mot.lower() in statuts_valides:
                nouveau_statut = mot.lower()
        if id_trouve and nouveau_statut:
            conn = sqlite3.connect(DB_PATH)
            c = conn.cursor()
            c.execute("UPDATE candidatures SET statut = ?, date_updated = date('now') WHERE id = ?", (nouveau_statut, id_trouve))
            conn.commit()
            conn.close()
            return f"✅ Candidature #{id_trouve} mise à jour : statut → {nouveau_statut}"
        return "Format : tracker statut [id] [statut]\nStatuts : envoyé, en attente, entretien, refus, accepté"

    # ── Ajouter une candidature
    else:
        # Extraire les infos du message
        urls = re.findall(r'https?://[^\s]+', message_complet)
        lien = urls[0] if urls else ""

        # Utiliser Qwen pour extraire entreprise et poste si URL présente
        if lien:
            try:
                headers = {"User-Agent": "Mozilla/5.0 Chrome/120.0.0.0"}
                r = requests.get(lien, headers=headers, timeout=10)
                soup = BeautifulSoup(r.text, "html.parser")
                for tag in soup(["script", "style"]): tag.decompose()
                texte = soup.get_text(separator="\n", strip=True)[:800]
                prompt = f"""Extrais le nom de l entreprise et le titre du poste.
Reponds UNIQUEMENT avec ce JSON : {{"entreprise": "...", "poste": "..."}}
Texte : {texte[:500]}"""
                rep = ollama.chat(model=MODELE, messages=[{"role": "user", "content": prompt}], options={"temperature": 0})
                txt = rep["message"]["content"].strip()
                debut = txt.find("{")
                fin = txt.rfind("}") + 1
                infos = json.loads(txt[debut:fin]) if debut != -1 else {}
                entreprise = infos.get("entreprise", "Inconnue")
                poste = infos.get("poste", "Inconnu")
            except:
                entreprise = "Inconnue"
                poste = "Inconnu"
        else:
            # Extraire depuis le message texte
            mots = message_complet.replace("tracker", "").replace("ajouter", "").strip().split()
            entreprise = mots[0] if len(mots) > 0 else "Inconnue"
            poste = " ".join(mots[1:]) if len(mots) > 1 else "Inconnu"

        # Calculer la date de relance (J+7 par défaut)
        date_auj = datetime.date.today()
        date_relance = (date_auj + datetime.timedelta(days=7)).isoformat()

        # Insérer dans la base (schéma server.py)
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute(
            "INSERT INTO candidatures (poste, entreprise, url, statut, status_code, date_next_action) VALUES (?, ?, ?, ?, ?, ?)",
            (poste, entreprise, lien, "envoyé", "applied", date_relance)
        )
        new_id = c.lastrowid
        conn.commit()
        conn.close()

        return (f"✅ Candidature #{new_id} ajoutée !\n\n"
                f"🏢 Entreprise : {entreprise}\n"
                f"💼 Poste : {poste}\n"
                f"📅 Envoyée le : {date_auj.strftime('%d/%m/%Y')}\n"
                f"🔔 Relance prévue : {(date_auj + datetime.timedelta(days=7)).strftime('%d/%m/%Y')}\n\n"
                f"Pour voir toutes tes candidatures : tracker voir\n"
                f"Pour mettre à jour : tracker statut {new_id} [statut]")



# ============================================================
# OUTIL 10 — GENERATEUR DE RAPPELS .ICS
# Génère un fichier .ics importable dans Google Calendar,
# Outlook ou Apple Calendar pour ne jamais oublier de relancer.
# Calcule automatiquement les dates optimales de relance.
# ============================================================

def generer_rappels_ics(message_complet):
    """
    Génère un fichier .ics avec les rappels de relance
    pour une candidature ou une liste de candidatures.
    """
    print(f"\n  [Outil 10] Génération des rappels .ics...")
    import datetime

    # Extraire les infos du message
    urls = re.findall(r'https?://[^\s]+', message_complet)
    entreprise = "Entreprise"
    poste = "Poste"
    date_candidature = datetime.date.today()

    if urls:
        try:
            headers = {"User-Agent": "Mozilla/5.0 Chrome/120.0.0.0"}
            r = requests.get(urls[0], headers=headers, timeout=10)
            soup = BeautifulSoup(r.text, "html.parser")
            for tag in soup(["script", "style"]): tag.decompose()
            texte = soup.get_text(separator="\n", strip=True)[:800]
            prompt = f"""Extrais l entreprise et le poste. JSON uniquement :
{{"entreprise": "...", "poste": "..."}}
Texte : {texte[:400]}"""
            rep = ollama.chat(model=MODELE, messages=[{"role": "user", "content": prompt}], options={"temperature": 0})
            txt = rep["message"]["content"].strip()
            debut = txt.find("{")
            fin = txt.rfind("}") + 1
            infos = json.loads(txt[debut:fin]) if debut != -1 else {}
            entreprise = infos.get("entreprise", "Entreprise")
            poste = infos.get("poste", "Poste")
        except:
            pass
    else:
        # Extraire depuis le message
        mots = message_complet.replace("rappel", "").replace("relance", "").strip().split()
        if mots:
            entreprise = mots[0]
        if len(mots) > 1:
            poste = " ".join(mots[1:3])

    # Calculer les dates de relance optimales
    # J+7 : première relance
    # J+14 : deuxième relance si pas de réponse
    # J+30 : clôture du dossier
    relance1 = date_candidature + datetime.timedelta(days=7)
    relance2 = date_candidature + datetime.timedelta(days=14)
    cloture = date_candidature + datetime.timedelta(days=30)

    def format_ics_date(d):
        return d.strftime("%Y%m%d")

    def format_uid():
        import uuid
        return str(uuid.uuid4())

    # Générer le contenu .ics
    ics_content = "BEGIN:VCALENDAR\r\n"
    ics_content += "VERSION:2.0\r\n"
    ics_content += "PRODID:-//Jobster//FR\r\n"
    ics_content += "CALSCALE:GREGORIAN\r\n"

    evenements = [
        (relance1, f"🔔 Relance J+7 — {entreprise}", f"Relance pour : {poste} chez {entreprise}\nSi pas de réponse, envoie un mail de suivi poli."),
        (relance2, f"🔔 Relance J+14 — {entreprise}", f"Deuxième relance pour : {poste} chez {entreprise}\nDernier rappel avant clôture."),
        (cloture, f"📁 Clôture dossier — {entreprise}", f"Clôture du suivi de candidature pour : {poste} chez {entreprise}\nMets à jour ton tracker Jobster."),
    ]

    for date_evt, titre, description in evenements:
        ics_content += "BEGIN:VEVENT\r\n"
        ics_content += f"UID:{format_uid()}\r\n"
        ics_content += f"DTSTART;VALUE=DATE:{format_ics_date(date_evt)}\r\n"
        ics_content += f"DTEND;VALUE=DATE:{format_ics_date(date_evt + datetime.timedelta(days=1))}\r\n"
        ics_content += f"SUMMARY:{titre}\r\n"
        ics_content += f"DESCRIPTION:{description}\r\n"
        ics_content += "BEGIN:VALARM\r\n"
        ics_content += "TRIGGER:-PT9H\r\n"
        ics_content += "ACTION:DISPLAY\r\n"
        ics_content += f"DESCRIPTION:Rappel : {titre}\r\n"
        ics_content += "END:VALARM\r\n"
        ics_content += "END:VEVENT\r\n"

    ics_content += "END:VCALENDAR\r\n"

    # Sauvegarder le fichier .ics
    nom_fichier = f"rappels_{entreprise.replace(' ', '_')}_{format_ics_date(date_candidature)}.ics"
    chemin_fichier = os.path.join(os.path.dirname(os.path.abspath(__file__)), nom_fichier)

    with open(chemin_fichier, "w", encoding="utf-8") as f:
        f.write(ics_content)

    print(f"  [Outil 10] Fichier .ics créé : {nom_fichier}")

    return (f"📅 RAPPELS DE RELANCE GÉNÉRÉS !\n\n"
            f"🏢 Entreprise : {entreprise}\n"
            f"💼 Poste : {poste}\n\n"
            f"📆 Dates de relance :\n"
            f"  🔔 J+7  : {relance1.strftime('%d/%m/%Y')} — Première relance\n"
            f"  🔔 J+14 : {relance2.strftime('%d/%m/%Y')} — Deuxième relance\n"
            f"  📁 J+30 : {cloture.strftime('%d/%m/%Y')} — Clôture du dossier\n\n"
            f"📁 Fichier .ics : {chemin_fichier}\n\n"
            f"💡 Importe ce fichier dans Google Calendar, Outlook ou Apple Calendar !")


def get_ft_token():
    """Récupère un token France Travail (réutilisable pour toutes les APIs)."""
    from dotenv import load_dotenv
    load_dotenv(dotenv_path=os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'backend', '.env'))
    url = "https://entreprise.francetravail.fr/connexion/oauth2/access_token"
    data = {
        "grant_type": "client_credentials",
        "client_id": os.getenv("FRANCE_TRAVAIL_CLIENT_ID"),
        "client_secret": os.getenv("FRANCE_TRAVAIL_CLIENT_SECRET"),
        "scope": "api_offresdemploiv2 nomenclatureRome o2dsoffre",
    }
    r = requests.post(url, params={"realm": "/partenaire"}, data=data)
    return r.json().get("access_token", "")


# ── API La Bonne Boite ────────────────────────────────────────
# Trouve les entreprises qui recrutent même sans offre publiée.
# Utilise l'API publique LBB (sans scope FT requis).
# ─────────────────────────────────────────────────────────────

def api_la_bonne_boite(metier, ville="Paris", latitude="48.8566", longitude="2.3522", distance=30):
    """
    Retourne les entreprises avec fort potentiel d'embauche via l'API publique LBB.
    """
    print(f"\n  [API La Bonne Boite] Recherche entreprises pour '{metier}' à {ville}...")
    try:
        lbb_headers = {
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36",
            "Referer": "https://labonneboite.francetravail.fr/",
            "Accept": "application/json",
        }

        # Résoudre le code ROME depuis les mots-clés métier
        # On prend 5 suggestions et on choisit celle dont le libellé contient le plus de mots du métier
        query = (metier or "développeur").lower()
        r_job = requests.get(
            "https://labonneboite.francetravail.fr/api/v2/autocomplete/jobs",
            params={"q": metier or "développeur", "limit": 5},
            headers=lbb_headers, timeout=10,
        )
        job_items = r_job.json().get("items", []) if r_job.status_code == 200 else []
        best = None
        best_score = -1
        for item in job_items:
            label = item.get("display", "").lower()
            score = sum(1 for w in query.split() if w in label)
            if score > best_score:
                best_score = score
                best = item
        if not best and job_items:
            best = job_items[0]
        rome_code = best["value"] if best else "M1805"
        job_label = best["display"] if best else (metier or "Développeur")

        # Résoudre le citycode INSEE depuis la ville
        r_loc = requests.get(
            "https://labonneboite.francetravail.fr/api/v2/autocomplete/location",
            params={"q": ville or "Paris", "limit": 1},
            headers=lbb_headers, timeout=10,
        )
        loc_items = r_loc.json().get("items", []) if r_loc.status_code == 200 else []
        citycode = loc_items[0]["value"] if loc_items else "75056"
        ville_label = loc_items[0]["display"] if loc_items else (ville or "Paris")

        # Recherche entreprises
        r = requests.get(
            "https://labonneboite.francetravail.fr/api/v2/search",
            params={
                "rome": rome_code,
                "citycode": citycode,
                "page": 1,
                "page_size": 10,
                "sort_by": "romes.hiring_potential",
                "sort_direction": "desc",
                "distance": distance,
            },
            headers=lbb_headers, timeout=15,
        )

        data = r.json()
        entreprises = data.get("items", [])
        total = data.get("hits", 0)

        if not entreprises:
            return (f"Aucune entreprise trouvée pour '{job_label}' à {ville_label}.\n\n"
                    f"💡 Essaie sur La Bonne Boîte : https://labonneboite.francetravail.fr/")

        resultat = f"🏢 **{total} entreprises** avec fort potentiel d'embauche pour **{job_label}** à **{ville_label}** :\n\n"
        for e in entreprises[:8]:
            nom = e.get("company_name") or e.get("office_name", "N/A")
            city = e.get("city", "N/A")
            dept = e.get("department", "")
            score = e.get("hiring_potential", 0)
            naf = e.get("naf_label", "")
            a_email = e.get("email") == "yes"

            resultat += f"• **{nom}** — {city} ({dept})\n"
            if isinstance(score, (int, float)) and score > 0:
                resultat += f"  🎯 Score recrutement : {score:.0f}/100"
            if naf:
                resultat += f" | {naf}"
            resultat += "\n"
            if a_email:
                resultat += "  ✉️ Candidature spontanée par email possible\n"
            resultat += "\n"

        resultat += f"👉 Voir toutes les entreprises : https://labonneboite.francetravail.fr/"
        return resultat

    except Exception as e:
        return f"Erreur La Bonne Boite : {e}"


# ── API Marché du Travail ─────────────────────────────────────
# Donne les statistiques d'un marché d'emploi :
# nombre d'offres, tension, salaire moyen, profils recrutés.
# ─────────────────────────────────────────────────────────────

def api_marche_travail(metier, departement=None):
    """
    Retourne les statistiques du marché du travail
    pour un métier dans un département.
    """
    print(f"\n  [API Marché du Travail] Stats pour '{metier}' dept {departement}...")
    try:
        token = get_ft_token()
        url = "https://api.francetravail.io/partenaire/marche-travail/v1/offres/statistiques"
        headers = {"Authorization": f"Bearer {token}", "Accept": "application/json"}
        params = {"motsCles": metier, "departement": departement}
        r = requests.get(url, headers=headers, params=params, timeout=15)
        data = r.json()

        # On demande à Qwen de synthétiser les données brutes
        prompt = f"""Voici des données brutes sur le marché du travail pour '{metier}' :
{json.dumps(data, ensure_ascii=False)[:1500]}

Fais un résumé clair en 5 points pour un candidat :
- Nombre d'offres disponibles
- Types de contrats dominants
- Niveau d'expérience demandé
- Salaires indicatifs si disponibles
- Conseil pour se positionner sur ce marché
"""
        reponse = ollama.chat(model=MODELE, messages=[{"role": "user", "content": prompt}])
        return "📊 MARCHÉ DU TRAVAIL :\n\n" + reponse["message"]["content"].strip()
    except Exception as e:
        return f"Erreur Marché du Travail : {e}"


# ── API Mes Évènements Emploi ─────────────────────────────────
# Trouve les salons, forums, ateliers emploi près de l'utilisateur.
# Utilise l'API publique mesevenementsemploi.francetravail.fr.
# ─────────────────────────────────────────────────────────────

def api_evenements_emploi(ville="Paris", departement=None):
    """
    Retourne les évènements emploi (salons, job datings, ateliers)
    à venir via l'API publique de mesevenementsemploi.francetravail.fr.
    """
    if departement is None:
        departement = ville_vers_dept(ville)
    print(f"\n  [API Évènements] Recherche évènements à {ville} (dept {departement})...")
    try:
        import time as _time
        mee_headers = {
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36",
            "Referer": "https://mesevenementsemploi.francetravail.fr/mes-evenements-emploi/evenements",
            "Accept": "application/json, text/plain, */*",
            "Content-Type": "application/json",
            "Origin": "https://mesevenementsemploi.francetravail.fr",
            "x-initialized-at": str(int(_time.time() * 1000)),
            "user_location": "2",
        }
        body = {
            "codeDepartement": departement or "",
            "codeRegion": "",
            "adresse": "",
            "modalitesAcces": [],
            "tagsOperation": [],
            "mustHavePlacesDisponibles": False,
            "tagsObjectif": [],
            "tagsPublic": [],
            "tagsType": [],
            "prerequis": [],
            "pageNumber": 0,
            "pageSize": 8,
            "sortBy": "date_evenement",
            "offset": 0,
        }
        r = requests.post(
            "https://mesevenementsemploi.francetravail.fr/api-candidat/mee/v1/de/evenement/all/filtered",
            json=body, headers=mee_headers, timeout=15,
        )
        if r.status_code != 200:
            return (f"🎪 Évènements emploi à {ville} :\n\n"
                    f"💡 Retrouve tous les salons et job datings près de chez toi sur :\n"
                    f"👉 https://mesevenementsemploi.francetravail.fr/\n\n"
                    f"Tu peux aussi appeler le **3949** (France Travail) pour les événements dans ta région.")

        data = r.json()
        evenements = data.get("items", [])
        total = data.get("count", 0)

        if not evenements:
            return (f"Aucun évènement emploi trouvé à {ville} pour le moment.\n\n"
                    f"💡 Essaie sur : https://mesevenementsemploi.francetravail.fr/")

        resultat = f"🎪 **{total} évènements emploi** dans la région **{ville}** :\n\n"
        for e in evenements[:8]:
            titre = e.get("titre", "N/A")
            date = e.get("dateEvenement", "N/A")
            heure_debut = e.get("heureDebut", "")
            heure_fin = e.get("heureFin", "")
            city = e.get("ville") or ""
            tag_type = (e.get("tagTypeEvenement") or {}).get("libelle", "")
            en_ligne = e.get("estEnLigne", False)
            preinscription = e.get("preinscription", False)
            nb_places = None
            for mode in (e.get("modeAcceesList") or []):
                nb_places = mode.get("nombrePlace")
                break

            resultat += f"• **{titre}**\n"
            horaire = f"{heure_debut}–{heure_fin}" if heure_debut else ""
            lieu_str = "🌐 En ligne" if en_ligne else (f"📍 {city}" if city else "")
            resultat += f"  📅 {date}"
            if horaire:
                resultat += f" {horaire}"
            if lieu_str:
                resultat += f" — {lieu_str}"
            resultat += "\n"
            if tag_type:
                resultat += f"  🏷️ {tag_type}"
            if nb_places:
                resultat += f" | {nb_places} places"
            if preinscription:
                resultat += " | Préinscription requise"
            resultat += "\n\n"

        resultat += f"👉 Voir tous les évènements : https://mesevenementsemploi.francetravail.fr/"
        return resultat

    except Exception as e:
        return f"Erreur Évènements emploi : {e}"


# ── API Référentiel des Agences ───────────────────────────────
# Localise les agences France Travail proches de l'utilisateur.
# Utile pour prendre rendez-vous avec un conseiller.
# ─────────────────────────────────────────────────────────────

def api_agences_france_travail(ville="Paris", departement=None):
    """
    Retourne les agences France Travail les plus proches
    dans un département donné.
    """
    if departement is None:
        departement = ville_vers_dept(ville)
    print(f"\n  [API Agences] Recherche agences à {ville} (dept {departement})...")
    try:
        token = get_ft_token()
        url = "https://api.francetravail.io/partenaire/referentiel-agences/v1/agences"
        headers = {"Authorization": f"Bearer {token}", "Accept": "application/json"}
        params = {"codeDepartement": departement}
        r = requests.get(url, headers=headers, params=params, timeout=15)
        agences = r.json()
        if not agences or not isinstance(agences, list):
            return f"Aucune agence trouvée à {ville}."
        resultat = f"🏢 Agences France Travail à {ville} :\n\n"
        for a in agences[:6]:
            resultat += f"• {a.get('libelleEtendu', a.get('libelle', 'N/A'))}\n"
            adresse = a.get('adresse', {})
            resultat += f"  📍 {adresse.get('ligne4', '')} {adresse.get('codePostal', '')} {adresse.get('localite', '')}\n"
            tel = a.get('telephone', '')
            if tel:
                resultat += f"  📞 {tel}\n"
            resultat += "\n"
        return resultat
    except Exception as e:
        return f"Erreur Référentiel Agences : {e}"


# ── API ROME 4.0 Métiers ──────────────────────────────────────
# Donne toutes les informations sur un métier :
# compétences requises, formations, salaires, évolutions.
# Utile pour s'orienter ou préparer une reconversion.
# ─────────────────────────────────────────────────────────────

def api_rome_metier(metier):
    """
    Retourne une fiche métier structurée via Ollama (fallback IA locale).
    L'API ROME 4.0 FT nécessite un scope non activé — on génère via Qwen directement.
    """
    print(f"\n  [API ROME] Fiche métier pour '{metier}' via Ollama...")
    try:
        prompt = f"""Tu es un conseiller en orientation professionnelle expert du marché du travail français.
Génère une fiche métier complète et précise pour : "{metier}"

Réponds EXACTEMENT avec ces 5 sections (conserve les titres exacts avec les ##) :

## Description
Décris ce que fait ce professionnel au quotidien en 2-3 phrases claires.

## Compétences clés
- liste de 5 à 7 compétences techniques et comportementales essentielles

## Formations recommandées
- liste de 3 à 5 diplômes ou formations adaptés (Bac+2, Bac+3, Master, Bootcamp...)

## Évolutions de carrière
- liste de 3 à 4 postes vers lesquels ce métier peut évoluer

## Conseil pour se lancer
1 ou 2 phrases de conseil pratique et actionnable pour débuter dans ce métier.

Réponds en français. Sois concis et précis."""

        reponse = ollama.chat(model=MODELE, messages=[{"role": "user", "content": prompt}])
        content = reponse["message"]["content"].strip()
        # Retirer les balises <think> si Qwen3 les génère
        import re as _re
        content = _re.sub(r'<think>.*?</think>', '', content, flags=_re.DOTALL).strip()
        return f"📚 FICHE MÉTIER — {metier.upper()} :\n\n{content}"
    except Exception as e:
        return f"Erreur ROME Métiers : {e}"

def comprendre_demande(message_utilisateur):
    """
    Utilise Qwen pour extraire le métier et la ville
    depuis le message de l'utilisateur.
    Retourne un dictionnaire avec 'keywords' et 'location'.
    """
    print("\n  [Agent] Analyse de votre demande...")

    prompt = f"""Tu es un assistant spécialisé dans la recherche d'emploi.
Analyse cette demande et extrait UNIQUEMENT le métier recherché et la ville.

Demande : "{message_utilisateur}"

Reponds UNIQUEMENT avec un JSON valide, sans aucun texte avant ou après :
{{"keywords": "le métier ou poste recherché", "location": "la ville mentionnée ou null si aucune ville n'est précisée"}}

Exemples :
- "je cherche un CDI chef de projet à Bordeaux" → {{"keywords": "chef de projet", "location": "Bordeaux"}}
- "stage développeur python" → {{"keywords": "développeur python", "location": null}}
- "alternance marketing digital Lyon" → {{"keywords": "marketing digital", "location": "Lyon"}}
- "ingénieur data science" → {{"keywords": "ingénieur data science", "location": null}}
- "commercial terrain Toulouse CDI" → {{"keywords": "commercial terrain", "location": "Toulouse"}}
- "find me a backend developer job in Nantes" → {{"keywords": "backend developer", "location": "Nantes"}}
- "software engineer remote" → {{"keywords": "software engineer", "location": null}}
"""

    try:
        reponse = ollama.chat(
            model=MODELE,
            messages=[{"role": "user", "content": prompt}],
            options={"temperature": 0}  # On veut une réponse précise, pas créative
        )
        texte = reponse["message"]["content"].strip()

        # On nettoie le texte pour extraire uniquement le JSON
        # Qwen peut ajouter du texte avant/après le JSON
        debut = texte.find("{")
        fin = texte.rfind("}") + 1
        if debut != -1 and fin > debut:
            json_texte = texte[debut:fin]
            resultat = json.loads(json_texte)
            # Normalise location : null / "null" / "" → None pour que server.py
            # utilise la localisation du profil plutôt que de forcer Paris.
            loc = resultat.get("location")
            if not loc or str(loc).lower() in ("null", "none", ""):
                resultat["location"] = None
            print(f"  [Agent] Métier détecté : {resultat.get('keywords', 'N/A')}")
            print(f"  [Agent] Ville détectée : {resultat.get('location', 'N/A (utilise profil)')}")
            return resultat
        else:
            # Fallback si le JSON n'est pas trouvé
            print("  [Agent] Impossible d'analyser la demande, utilisation des valeurs par défaut")
            return {"keywords": message_utilisateur, "location": None}
    except Exception as e:
        print(f"  [Agent] Erreur d'analyse : {e}")
        return {"keywords": message_utilisateur, "location": None}


# ============================================================
# ETAPE 2 — LANCER LE SCRAPER
# Une fois qu'on sait quoi chercher et où, on lance
# le scraper pour aller chercher les vraies offres.
# ============================================================

def _normaliser(texte):
    """Minuscule, sans accents/espaces multiples — pour comparer deux offres."""
    texte = (texte or "").lower().strip()
    texte = re.sub(r'[^a-z0-9]+', ' ', texte)
    return re.sub(r'\s+', ' ', texte).strip()


def dedupliquer_offres(offres):
    """
    Supprime les doublons entre sources (ex: même offre sur France Travail
    ET relayée par Adzuna). Deux offres sont considérées identiques si leur
    titre + entreprise normalisés correspondent. Garde la première occurrence
    (priorité aux sources listées en premier dans aggregate_all).
    """
    vus = set()
    resultat = []
    for o in offres:
        cle = (_normaliser(o.get("titre")), _normaliser(o.get("entreprise")))
        if cle in vus:
            continue
        vus.add(cle)
        resultat.append(o)
    if len(resultat) < len(offres):
        print(f"  [Scraper] Dédoublonnage : {len(offres)} -> {len(resultat)} offres")
    return resultat


def trier_par_pertinence(offres, keywords):
    """
    Trie les offres par pertinence par rapport aux mots-clés recherchés.
    Score = nombre de mots-clés présents dans le titre (le plus pertinent
    en premier). Les offres à score égal gardent leur ordre d'origine
    (tri stable) pour ne pas mélanger les sources entre elles.
    """
    mots = [m for m in _normaliser(keywords).split(" ") if m]
    if not mots:
        return offres

    def score(offre):
        titre_norm = _normaliser(offre.get("titre"))
        return -sum(1 for m in mots if m in titre_norm)

    return sorted(offres, key=score)


def lancer_scraper(keywords, location):
    """
    Lance les sources rapides (APIs) + scrapers Playwright en parallèle.

    Pipeline :
    • France Travail + Adzuna → répondent en < 5 s (API officielle)
    • Indeed + WTTJ (Playwright) → lancés en parallèle dans des threads démons
      avec un timeout de 25 s. Si Chromium n'est pas installé ou si la page
      bloque, le thread se termine silencieusement et on garde les offres API.
    • Les résultats fusionnés sont dédoublonnés puis triés par pertinence
      avant d'être retournés à l'agent.
    """
    import threading

    print(f"\n  [Scraper] Recherche de '{keywords}' à '{location}'...")
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

    try:
        from jobster_scraper import (
            scrape_france_travail, scrape_adzuna,
            scrape_indeed, scrape_wttj,
        )
    except ImportError as e:
        print(f"  [Scraper] Import error : {e}")
        return []

    offres = []
    dept = ville_vers_dept(location)

    # ── Sources rapides (synchrones) ──────────────────────────────────
    try:
        if dept:
            offres += scrape_france_travail(keywords, dept)
        else:
            offres += scrape_france_travail(keywords)
    except Exception as e:
        print(f"  [Scraper] France Travail erreur : {e}")

    try:
        offres += scrape_adzuna(keywords, location)
    except Exception as e:
        print(f"  [Scraper] Adzuna erreur : {e}")

    # ── Playwright en parallèle (threads démons, timeout 25 s) ───────
    playwright_results = []
    lock = threading.Lock()

    def run_indeed():
        try:
            result = scrape_indeed(keywords, location, max_results=5)
            with lock:
                playwright_results.extend(result)
        except Exception as e:
            print(f"  [Scraper] Indeed thread erreur : {e}")

    def run_wttj():
        try:
            result = scrape_wttj(keywords, location, max_results=5)
            with lock:
                playwright_results.extend(result)
        except Exception as e:
            print(f"  [Scraper] WTTJ thread erreur : {e}")

    t_indeed = threading.Thread(target=run_indeed, daemon=True)
    t_wttj   = threading.Thread(target=run_wttj,   daemon=True)
    t_indeed.start()
    t_wttj.start()
    t_indeed.join(timeout=25)
    t_wttj.join(timeout=25)

    offres += playwright_results

    offres = dedupliquer_offres(offres)
    offres = trier_par_pertinence(offres, keywords)

    print(f"  [Scraper] {len(offres)} offres au total "
          f"(dont {len(playwright_results)} Playwright)")
    return offres


# ============================================================
# ETAPE 3 — LIRE LES OFFRES ET REDIGER LA REPONSE
# L'agent lit toutes les offres trouvées et rédige une
# réponse claire et utile pour l'utilisateur.
# ============================================================

def rediger_reponse(offres, demande_originale):
    """
    Utilise Qwen pour analyser les offres et rédiger
    une réponse personnalisée pour l'utilisateur.
    """
    if not offres:
        return "Je n'ai trouvé aucune offre correspondant à votre recherche. Essayez avec d'autres mots-clés ou une autre ville."

    print(f"\n  [Agent] Analyse de {len(offres)} offres et rédaction de la réponse...")

    # On prépare un résumé des offres pour Qwen
    # On limite à 15 offres pour ne pas dépasser la mémoire du modèle
    offres_texte = ""
    for i, offre in enumerate(offres[:15], 1):
        offres_texte += f"{i}. {offre['titre']} — {offre['entreprise']} — {offre['lieu']} ({offre['source']})\n"

    prompt = f"""Tu es Jobster, un assistant IA sympathique spécialisé dans la recherche d'emploi.

L'utilisateur a demandé : "{demande_originale}"

Voici les offres trouvées :
{offres_texte}

Rédige une réponse en français qui :
1. Annonce combien d'offres ont été trouvées
2. Présente les 3 meilleures offres de façon claire avec le titre, l'entreprise et la source
3. Donne un conseil personnalisé pour postuler
4. Reste enthousiaste et encourageant

Réponds directement sans introduction inutile.
"""

    try:
        reponse = ollama.chat(
            model=MODELE,
            messages=[{"role": "user", "content": prompt}],
            options={"temperature": 0.7}  # Un peu de créativité pour la réponse
        )
        return reponse["message"]["content"].strip()
    except Exception as e:
        # Si Qwen échoue, on génère une réponse simple sans IA
        print(f"  [Agent] Erreur de génération : {e}")
        reponse_simple = f"J'ai trouvé {len(offres)} offres pour vous !\n\n"
        for i, offre in enumerate(offres[:5], 1):
            reponse_simple += f"{i}. {offre['titre']} — {offre['entreprise']} ({offre['source']})\n"
        return reponse_simple


# ============================================================
# BOUCLE PRINCIPALE
# C'est ici que l'agent attend les questions de l'utilisateur
# et y répond en boucle jusqu'à ce qu'il tape "quitter".
# ============================================================

def lancer_agent():
    """Lance l'agent en mode conversationnel."""
    print("=" * 60)
    print("  JOBSTER — Agent IA de recherche d'emploi")
    print("  Propulsé par Ollama + Qwen")
    print("  Tapez 'quitter' pour arrêter")
    print("=" * 60)
    print()
    print("  Bonjour ! Je suis Jobster, votre assistant emploi.")
    print()
    print("  Dites-moi ce que vous cherchez et je m'occupe du reste.")
    print()
    print("  Exemples :")
    print("  → 'chef de projet à Lyon'")
    print("  → 'analyse cette offre https://...'")
    print("  → 'tracker voir'  pour voir vos candidatures")
    print()

    while True:
        # ── On attend la question de l'utilisateur
        try:
            demande = input("Vous : ").strip()
        except KeyboardInterrupt:
            print("\n\n  Au revoir !")
            break

        if not demande:
            continue

        if demande.lower() in ["quitter", "quit", "exit", "bye", "au revoir"]:
            print("\n  Jobster : Au revoir et bonne chance dans votre recherche !")
            break

        print()

        # ── Détection par mots-clés explicites
        urls_trouvees = re.findall(r'https?://[^\s]+', demande)
        chemins_pdf = re.findall(r'[A-Za-z]:\\[^\s]+\.pdf|/[^\s]+\.pdf', demande, re.IGNORECASE)
        d = demande.lower()

        # Mots-clés par outil — chaque outil a ses déclencheurs uniques
        mots_lettre    = ["lettre", "motivation", "rédige lettre", "génère lettre"]
        mots_cv        = ["adapter mon cv", "adapte mon cv", "cv adapté", "génère mon cv", "cv personnalisé", "optimise mon cv"]
        mots_matching  = ["match", "score", "compatible", "correspond", "convient", "évalue profil", "analyse profil"]
        mots_analyse   = ["analyse cette offre", "décrypte offre", "explique offre", "que dit offre", "détail offre"]
        mots_copier    = ["prépare mail", "mail candidature", "texte formulaire", "copier coller candidature", "mail postuler"]
        mots_urls      = ["liens utiles", "génère liens", "trouve recruteur", "liens candidature", "ouvre glassdoor"]
        mots_rapport   = ["rapport entreprise", "infos entreprise", "recherche entreprise", "que sais-tu de", "avis employés", "analyse entreprise", "glassdoor"]
        mots_tracker   = ["tracker", "mes candidatures", "ajouter candidature", "voir candidatures", "suivi candidature", "statut candidature"]
        mots_rappel    = ["rappel", "relance", "calendrier candidature", "date relance", "ics", "agenda"]
        # APIs France Travail secondaires
        mots_bonne_boite = ["bonne boite", "entreprises qui recrutent", "candidature spontanée", "quelles entreprises recrutent"]
        mots_marche    = ["marché du travail", "statistiques emploi", "tension emploi", "comment se porte le marché"]
        mots_evenement = ["événement emploi", "salon emploi", "forum emploi", "job dating", "atelier emploi"]
        mots_agence    = ["agence france travail", "agence pôle emploi", "conseiller france travail", "agence ft"]
        mots_rome      = ["fiche métier", "compétences pour devenir", "reconversion vers", "rome", "formation pour être"]

        # ── OUTIL 4 : Lettre de motivation
        if any(m in d for m in mots_lettre) and urls_trouvees:
            print(f"  [Agent] → Outil 4 : lettre de motivation")
            reponse = generer_lettre_motivation(demande)

        # ── OUTIL 5 : CV adapté en PDF
        elif any(m in d for m in mots_cv) and urls_trouvees:
            print(f"  [Agent] → Outil 5 : CV adapté en PDF")
            reponse = generer_cv_adapte(demande)

        # ── OUTIL 6 : Mail + texte formulaire + LinkedIn
        elif any(m in d for m in mots_copier) and urls_trouvees:
            print(f"  [Agent] → Outil 6 : génération textes candidature")
            reponse = generer_copier_coller(demande)

        # ── OUTIL 7 : URLs intelligentes (LinkedIn recruteur, Glassdoor...)
        elif any(m in d for m in mots_urls) and urls_trouvees:
            print(f"  [Agent] → Outil 7 : génération URLs intelligentes")
            reponse = generer_urls_intelligentes(urls_trouvees[0])

        # ── OUTIL 8 : Rapport entreprise (Glassdoor + News + Société.com)
        elif any(m in d for m in mots_rapport):
            print(f"  [Agent] → Outil 8 : rapport entreprise")
            reponse = scraper_infos_entreprise(demande)

        # ── OUTIL 9 : Tracker de candidatures SQLite
        elif any(m in d for m in mots_tracker):
            print(f"  [Agent] → Outil 9 : tracker candidatures")
            reponse = tracker_candidature(demande)

        # ── OUTIL 10 : Rappels .ics
        elif any(m in d for m in mots_rappel):
            print(f"  [Agent] → Outil 10 : rappels .ics")
            reponse = generer_rappels_ics(demande)

        # ── OUTIL 3 : Matching CV/offre
        elif any(m in d for m in mots_matching) and urls_trouvees:
            print(f"  [Agent] → Outil 3 : matching CV/offre")
            reponse = calculer_matching(demande)

        # ── OUTIL 2 : Analyse d'offre depuis URL
        elif any(m in d for m in mots_analyse) and urls_trouvees:
            print(f"  [Agent] → Outil 2 : analyse de l'offre")
            reponse = analyser_offre_url(urls_trouvees[0])

        # ── URL seule courte → analyse automatique
        elif urls_trouvees and len(demande.strip()) < 80:
            print(f"  [Agent] → Outil 2 : analyse de l'offre (URL seule)")
            reponse = analyser_offre_url(urls_trouvees[0])

        # ── API La Bonne Boite
        elif any(m in d for m in mots_bonne_boite):
            print(f"  [Agent] → API La Bonne Boite")
            infos = comprendre_demande(demande)
            reponse = api_la_bonne_boite(infos.get("keywords", ""))

        # ── API Marché du Travail
        elif any(m in d for m in mots_marche):
            print(f"  [Agent] → API Marché du Travail")
            infos = comprendre_demande(demande)
            reponse = api_marche_travail(infos.get("keywords", demande))

        # ── API Evènements Emploi
        elif any(m in d for m in mots_evenement):
            print(f"  [Agent] → API Evènements Emploi")
            infos = comprendre_demande(demande)
            loc = infos.get("location", "Paris")
            reponse = api_evenements_emploi(loc, ville_vers_dept(loc))

        # ── API Agences France Travail
        elif any(m in d for m in mots_agence):
            print(f"  [Agent] → API Agences France Travail")
            infos = comprendre_demande(demande)
            loc = infos.get("location", "Paris")
            reponse = api_agences_france_travail(loc, ville_vers_dept(loc))

        # ── API ROME Métiers
        elif any(m in d for m in mots_rome):
            print(f"  [Agent] → API ROME Métiers")
            infos = comprendre_demande(demande)
            reponse = api_rome_metier(infos.get("keywords", demande))

        # ── OUTIL 1 : Recherche d'offres (cas par défaut)
        else:
            print(f"  [Agent] → Outil 1 : recherche d'offres")
            infos = comprendre_demande(demande)
            keywords = infos.get("keywords", demande)
            location = infos.get("location", "Paris")
            offres = lancer_scraper(keywords, location)
            reponse = rediger_reponse(offres, demande)

        print()
        print("  Jobster :", reponse)
        print()
        print("-" * 60)
        print()


# ============================================================
# BOUTON ON — Lance l'agent quand on tape :
# python jobster_agent.py
# ============================================================

if __name__ == "__main__":
    lancer_agent()


def ville_vers_dept(ville):
    table = {
        "paris": "75", "lyon": "69", "marseille": "13", "toulouse": "31",
        "bordeaux": "33", "lille": "59", "nantes": "44", "strasbourg": "67",
        "nice": "06", "rennes": "35", "montpellier": "34", "grenoble": "38",
        "rouen": "76", "toulon": "83", "dijon": "21", "angers": "49",
        "nimes": "30", "aix-en-provence": "13", "brest": "29", "limoges": "87",
        "clermont-ferrand": "63", "amiens": "80", "metz": "57", "nancy": "54",
        "perpignan": "66", "pau": "64", "caen": "14", "reims": "51",
        "le havre": "76", "saint-etienne": "42",
        # Recherche nationale — chaîne vide = pas de filtre département
        "france": "", "national": "", "": "",
    }
    return table.get(ville.lower().strip(), "75")
